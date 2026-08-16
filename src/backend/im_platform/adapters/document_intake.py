"""Document intake assistant for building the ab initio structural Investment
Register from primary deal documents.

This module does NOT attempt to auto-populate legal/financial facts. It:
1. Scans each portfolio company folder for likely primary documents (term
   sheets, investment summaries, purchase/subscription agreements, SAFEs).
2. Extracts raw text from the best candidates so a human can read and
   transcribe the real figures.
3. Produces a draft Investment Register sheet with blank canonical columns
   for a human to fill in and confirm, linked to the source document used.

Human confirmation is required before any row is treated as truth -- this
matches the platform's governance principle that state-changing facts need
human sign-off.
"""

from __future__ import annotations

from dataclasses import dataclass
from pathlib import Path

import pandas as pd
from openpyxl.cell.cell import ILLEGAL_CHARACTERS_RE

# Ordered by how reliably a document of this kind states final deal terms.
KEYWORD_PRIORITY = [
    "investment summary",
    "term sheet",
    "note purchase",
    "purchase and sale agreement",
    "subscription agreement",
    "share subscription",
    "shareholders agreement",
    "safe",
    "side letter",
]

DOCUMENT_EXTENSIONS = {".pdf", ".docx"}
IMAGE_EXTENSIONS = {".jpg", ".jpeg", ".png"}


@dataclass
class CandidateDocument:
    company_folder: str
    file_path: Path
    matched_keyword: str
    priority_rank: int


def discover_company_folders(investments_root: Path) -> list[Path]:
    """List immediate portfolio-company subfolders under the Equity root."""
    equity_root = investments_root / "0. E Q U I T Y"
    return sorted(p for p in equity_root.iterdir() if p.is_dir())


def discover_fund_folders(investments_root: Path) -> list[Path]:
    """List immediate fund-vehicle subfolders under the Fund Investment root."""
    fund_root = investments_root / "1. F U N D - I N V E S T M E N T"
    if not fund_root.exists():
        return []
    return sorted(p for p in fund_root.iterdir() if p.is_dir())


# Filenames containing these are per-entity legal/financial statements, so the
# text before the separator is a real sub-vehicle legal name -- useful for
# fund-of-funds structures (e.g. MGX) that have several distinct LP entities
# nested under one umbrella fund folder.
SUBVEHICLE_STATEMENT_KEYWORDS = [
    "capital account",
    "quarterly report",
    "audited financial statement",
    "limited partnership agreement",
    "side letter",
]

STATEMENT_EXTENSIONS = {".pdf", ".docx", ".xlsx"}


def discover_fund_subvehicles(fund_folder: Path) -> list[str]:
    """Derive distinct legal sub-vehicle names from statement filenames inside
    a fund folder."""
    names: set[str] = set()
    for path in fund_folder.rglob("*"):
        if not path.is_file() or path.suffix.lower() not in STATEMENT_EXTENSIONS:
            continue
        stem = path.stem
        lowered = stem.lower()
        if not any(keyword in lowered for keyword in SUBVEHICLE_STATEMENT_KEYWORDS):
            continue
        positions = [pos for pos in (stem.find(" - "), stem.find("~")) if pos != -1]
        split_pos = min(positions) if positions else -1
        candidate = stem[:split_pos].strip() if split_pos != -1 else stem.strip()
        if candidate:
            names.add(candidate)
    return sorted(names)


def _match_keyword(file_stem: str) -> tuple[str, int] | None:
    lowered = file_stem.lower()
    for rank, keyword in enumerate(KEYWORD_PRIORITY):
        if keyword in lowered:
            return keyword, rank
    return None


def find_candidate_documents(company_folder: Path) -> list[CandidateDocument]:
    candidates: list[CandidateDocument] = []
    for path in company_folder.rglob("*"):
        if not path.is_file() or path.suffix.lower() not in DOCUMENT_EXTENSIONS:
            continue
        match = _match_keyword(path.stem)
        if match is None:
            continue
        keyword, rank = match
        candidates.append(CandidateDocument(company_folder.name, path, keyword, rank))
    return sorted(candidates, key=lambda c: c.priority_rank)


def find_fallback_documents(company_folder: Path) -> list[CandidateDocument]:
    """When no keyword-matched document exists, surface any doc/pdf/image file
    found so a human still has something to open (e.g. a scanned promissory
    note saved only as an image)."""
    fallback: list[tuple[int, Path]] = []
    for path in company_folder.rglob("*"):
        if not path.is_file():
            continue
        suffix = path.suffix.lower()
        if suffix not in DOCUMENT_EXTENSIONS and suffix not in IMAGE_EXTENSIONS:
            continue
        depth = len(path.relative_to(company_folder).parts)
        fallback.append((depth, path))
    fallback.sort(key=lambda item: item[0])
    return [
        CandidateDocument(company_folder.name, path, "(fallback - no keyword match)", 999)
        for _, path in fallback[:1]
    ]


def extract_docx_text(path: Path) -> str:
    import docx

    document = docx.Document(str(path))
    parts = [p.text for p in document.paragraphs if p.text.strip()]
    for table in document.tables:
        for row in table.rows:
            parts.append(" | ".join(cell.text.strip() for cell in row.cells))
    return "\n".join(parts)


def extract_pdf_text(path: Path, max_pages: int = 6) -> str:
    import pdfplumber

    parts = []
    with pdfplumber.open(str(path)) as pdf:
        for page in pdf.pages[:max_pages]:
            text = page.extract_text() or ""
            if not text.strip():
                text = _ocr_pdf_page(page)
            if text.strip():
                parts.append(text)
    return "\n".join(parts)


def _configure_tesseract() -> None:
    import shutil

    import pytesseract

    if shutil.which("tesseract"):
        return
    candidate = Path.home() / "AppData" / "Local" / "Programs" / "Tesseract-OCR" / "tesseract.exe"
    if candidate.exists():
        pytesseract.pytesseract.tesseract_cmd = str(candidate)


def _ocr_pdf_page(page) -> str:
    import pytesseract

    _configure_tesseract()
    try:
        image = page.to_image(resolution=200).original
        return pytesseract.image_to_string(image)
    except Exception as exc:
        return f"[OCR FAILED: {exc}]"


def extract_image_text(path: Path) -> str:
    import pytesseract
    from PIL import Image

    _configure_tesseract()
    with Image.open(path) as image:
        return pytesseract.image_to_string(image)


def extract_document_text(path: Path) -> str:
    suffix = path.suffix.lower()
    try:
        if suffix == ".docx":
            text = extract_docx_text(path)
        elif suffix == ".pdf":
            text = extract_pdf_text(path)
        elif suffix in IMAGE_EXTENSIONS:
            text = extract_image_text(path)
        else:
            return "[Unsupported file type for text extraction]"
    except Exception as exc:  # a single bad file should not abort the whole scan
        return f"[EXTRACTION FAILED: {exc}]"
    return ILLEGAL_CHARACTERS_RE.sub("", text)


REGISTER_DRAFT_COLUMNS = [
    "investment_id",
    "entity_id",
    "fund_vehicle_id",
    "instrument_type",
    "initial_commitment_amount",
    "investment_currency",
    "close_date",
    "lifecycle_state",
    "lifecycle_state_date",
    "source_document",
    "confirmed_by",
    "confirmed_date",
]


def build_intake_workbook(investments_root: Path, output_path: Path, top_n_per_company: int = 2) -> Path:
    extract_rows = []
    draft_rows = []

    for company_folder in discover_company_folders(investments_root):
        candidates = find_candidate_documents(company_folder)[:top_n_per_company]

        if not candidates:
            candidates = find_fallback_documents(company_folder)

        if not candidates:
            extract_rows.append(
                {
                    "company_folder": company_folder.name,
                    "file_path": "",
                    "matched_keyword": "",
                    "extracted_text": "NO CANDIDATE DOCUMENT FOUND - needs manual review",
                }
            )
            draft_row = {col: "" for col in REGISTER_DRAFT_COLUMNS}
            draft_row["entity_id"] = company_folder.name
            draft_rows.append(draft_row)
            continue

        for candidate in candidates:
            text = extract_document_text(candidate.file_path)
            extract_rows.append(
                {
                    "company_folder": candidate.company_folder,
                    "file_path": str(candidate.file_path),
                    "matched_keyword": candidate.matched_keyword,
                    "extracted_text": text[:32000],
                }
            )

        draft_row = {col: "" for col in REGISTER_DRAFT_COLUMNS}
        draft_row["entity_id"] = company_folder.name
        draft_row["source_document"] = str(candidates[0].file_path)
        draft_rows.append(draft_row)

    extract_df = pd.DataFrame(extract_rows)
    draft_df = pd.DataFrame(draft_rows, columns=REGISTER_DRAFT_COLUMNS)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    with pd.ExcelWriter(output_path, engine="openpyxl") as writer:
        draft_df.to_excel(writer, index=False, sheet_name="Investment_Register_Draft")
        extract_df.to_excel(writer, index=False, sheet_name="Source_Document_Extract")

    return output_path


def append_draft_rows(intake_path: Path, new_rows: pd.DataFrame) -> Path:
    """Add rows to an existing intake workbook's draft register without
    disturbing the Source_Document_Extract sheet, skipping entity_ids already
    present so re-runs don't create duplicates."""
    existing_draft = pd.read_excel(intake_path, sheet_name="Investment_Register_Draft").fillna("")
    existing_extract = pd.read_excel(intake_path, sheet_name="Source_Document_Extract").fillna("")

    already_present = set(existing_draft["entity_id"])
    to_add = new_rows[~new_rows["entity_id"].isin(already_present)]

    combined = pd.concat([existing_draft, to_add], ignore_index=True)

    with pd.ExcelWriter(intake_path, engine="openpyxl") as writer:
        combined.to_excel(writer, index=False, sheet_name="Investment_Register_Draft")
        existing_extract.to_excel(writer, index=False, sheet_name="Source_Document_Extract")

    return intake_path


# Markers written into the "confirmed_by" note whenever a row relies on
# something other than a confirmed primary source document (a gap, a tracker
# fallback, a deferred question, or a follow-up still owed). Kept as a plain
# list so it's easy to extend as new phrasing patterns get used.
SOURCE_GAP_MARKERS = [
    "GAP:",
    "NEEDS HUMAN INPUT",
    "STILL DEFERRED",
    "STILL NEEDED",
    "TRACKER FALLBACK",
    "TRACKER ONLY",
    "NEEDS FOLLOW-UP",
    "DOCUMENT STILL NEEDED",
    "PARKED",
]


def refresh_needs_source_documents(intake_path: Path) -> Path:
    """Rebuild the 'Needs_Source_Documents' tracking sheet by scanning the
    draft register's confirmed_by notes for known gap markers. Safe to call
    repeatedly - it always reflects the current state of the register rather
    than being manually maintained."""
    draft = pd.read_excel(intake_path, sheet_name="Investment_Register_Draft").fillna("")
    extract = pd.read_excel(intake_path, sheet_name="Source_Document_Extract").fillna("")

    def _matched_markers(note: str) -> str:
        return ", ".join(m.rstrip(":") for m in SOURCE_GAP_MARKERS if m in note)

    rows = []
    for _, row in draft.iterrows():
        note = str(row.get("confirmed_by", ""))
        markers = _matched_markers(note)
        if not markers:
            continue
        rows.append(
            {
                "entity_id": row.get("entity_id", ""),
                "investment_id": row.get("investment_id", ""),
                "gap_type": markers,
                "note": note,
            }
        )

    gaps_df = pd.DataFrame(rows, columns=["entity_id", "investment_id", "gap_type", "note"])

    with pd.ExcelWriter(intake_path, engine="openpyxl") as writer:
        draft.to_excel(writer, index=False, sheet_name="Investment_Register_Draft")
        extract.to_excel(writer, index=False, sheet_name="Source_Document_Extract")
        gaps_df.to_excel(writer, index=False, sheet_name="Needs_Source_Documents")

    return intake_path
